import os
from docx import Document
from fpdf import FPDF
from config import WORD_FOLDER, PDF_FOLDER


def save_word(user_id, name, orig, trans):
    folder = os.path.join(WORD_FOLDER, str(user_id))
    os.makedirs(folder, exist_ok=True)
    path = os.path.join(folder, f"{name}.docx")
    doc = Document()
    doc.add_heading('Распознанный текст', level=1)
    doc.add_paragraph(orig)
    doc.add_heading('Переведённый текст', level=1)
    doc.add_paragraph(trans)
    doc.save(path)
    return path


def save_pdf(user_id, name, orig, trans):
    folder = os.path.join(PDF_FOLDER, str(user_id))
    os.makedirs(folder, exist_ok=True)
    path = os.path.join(folder, f"{name}.pdf")
    pdf = FPDF()
    pdf.add_page()
    pdf.add_font('DejaVu', '', 'DejaVuSans.ttf', uni=True)
    pdf.set_font('DejaVu', size=12)
    pdf.multi_cell(0, 10, f"Распознанный текст:\n{orig}\n\nПеревод:\n{trans}")
    pdf.output(path)
    return path


